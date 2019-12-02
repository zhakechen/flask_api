from docx import Document
from flask import Blueprint, render_template, redirect, url_for, jsonify, request
from flask_login import login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash
from datetime import datetime
import json

from application.users.models import *
from application import db
from application.projects.forms import *

# 生成蓝图对象
projects = Blueprint('projects', __name__)

# 删除项目
@projects.route('/delete_project', methods=['POST'])
@login_required
def delete_project():
    json_str = request.get_data(as_text=True)
    json_data = json.loads(json_str)
    project_id = json_data['project_id']
    project = Project.query.get(int(project_id))
    db.session.delete(project)
    db.session.commit()
    return jsonify({'code': 200, 'message': '成功', 'project_id': project_id})

# 项目详情
@projects.route('/detail/<int:project_id>', methods=['GET'])
@login_required
def detail(project_id):
    project = Project.query.get(int(project_id))
    modules = project.modules
    return render_template('projects/detail.html', project=project, modules=modules)    

# 模块详情
@projects.route('/modules/<int:module_id>', methods=['GET'])
@login_required
def modules(module_id):
    module = Module.query.get(int(module_id))
    interfaces = module.interfaces
    return render_template('projects/module_detail.html', module=module, interfaces=interfaces)    

# 我的项目列表
@projects.route('/', methods=['GET', 'POST'])
@login_required
def index():
    projects = Project.query.filter_by(creator=current_user).all()
    # if request.method == 'POST':
    #     json_str = request.get_data(as_text=True)
    #     json_data = json.loads(json_str)
        # key = json_data['key']
        # result = Project.query.filter(Project.name.like('%'+key+'%')).first().name

        # return jsonify({'code': 200, 'message': '成功', 'result': result})

    key = request.args.get('key')
    if key:
        projects = Project.query.filter(Project.name.like('%'+key+'%')).all()
    return render_template('projects/index.html', projects=projects)

# 创建项目
@projects.route('/create_project', methods=['GET', 'POST'])
@login_required
def create_project():
    form = ProjectForm()
    if form.validate_on_submit():
        project = Project()
        project.name = form.project_name.data
        project.info = form.project_info.data
        project.creator = current_user
        project.created_time = datetime.now()

        project.updator = current_user
        project.updated_time = datetime.now()

        db.session.add(project)
        db.session.commit()
        return redirect(url_for('projects.index'))

    return render_template('projects/create_project.html', form=form)


# 自定义搜索




# 创建模块
@projects.route('/create_module/<int:project_id>', methods=['GET', 'POST'])
@login_required
def create_module(project_id):
    form = ModuleForm()
    project = Project.query.get(int(project_id))

    if form.validate_on_submit():
        module = Module()
        module.name = form.module_name.data
        module.info = form.module_info.data
        module.creator = current_user
        module.created_time = datetime.now()
        module.project = project
        module.updator = current_user
        module.updated_time = datetime.now()

        db.session.add(module)
        db.session.commit()
        return redirect(url_for('projects.detail', project_id=project.id))

    return render_template('projects/create_module.html', form=form, project=project)    

# 创建接口
@projects.route('/create_interface/<int:module_id>', methods=['GET', 'POST'])
@login_required
def create_interface(module_id):
    form = InterfaceForm()
    module = Module.query.get(int(module_id))

    if form.validate_on_submit():
        interface = Interface()
        interface.name = form.name.data
        interface.info = form.info.data
        interface.url = form.url.data
        interface.method_type = form.method_type.data
        
        interface.creator = current_user
        interface.created_time = datetime.now()
        interface.module = module
        interface.project_id = module.project_id
        interface.updator = current_user
        interface.updated_time = datetime.now()

        db.session.add(module)
        db.session.commit()
        return redirect(url_for('projects.modules', module_id=module.id))

    return render_template('projects/create_interface.html', form=form, module=module)    

# 接口修改
@projects.route('/edit_interface/<int:interface_id>', methods=['GET', 'POST'])
@login_required
def edit_interface(interface_id):
    interface = Interface.query.get(int(interface_id))
    if request.method == 'POST':    
        json_str = request.get_data(as_text=True)
        json_data = json.loads(json_str)
        request_data = json_data['request_data']
        response_data = json_data['response_data']

        interface.request_body = json.dumps(request_data, ensure_ascii=False)
        interface.response_body = json.dumps(response_data, ensure_ascii=False)

        interface.updator = current_user
        interface.updated_time = datetime.now()

        db.session.add(interface)
        db.session.commit()
        return jsonify({'code': 200, 'message': '成功', 'interface_id': interface.id})

    return render_template('projects/edit_interface.html', interface=interface)    

# 获取接口详情
@projects.route('/interfaces/<int:interface_id>', methods=['GET'])
@login_required
def interfaces(interface_id):
    interface = Interface.query.get(int(interface_id))
    json_data = interface.to_dict()
    return jsonify({'code': 200, 'message': '成功', 'data': json_data})



# 导出word文档
@projects.route('/export_word/<int:project_id>', methods=['GET'])
@login_required
def export_word(project_id):
    interfaces = Project.query.get(int(project_id))
    doc = Document()
    # paragraph = doc.add_paragraph(project.name)
    # modules = project.modules
    # for module in modules:
    #     doc.add_paragraph(module.name)
    #     for interface in module.interfaces:
    #         doc.add_paragraph(interface.name)

    doc.save('a.docx')

    return jsonify({'code': 200, 'message': '成功'})

# 在线预览
@projects.route('/onlineread/', methods=['GET'])
@login_required
def onlineread():
    with open('doc/abc.pdf', 'rb') as f:
        content = f.read()
    response = make_response(content)
    response.headers['Content-Type'] = 'application/pdf'

    return response





